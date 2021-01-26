#get_file_word.py
#!/usr/bin/env python
# coding: utf-8

# In[ ]:


# Importar bibliotecas
from os import listdir
from os.path import isfile, join
import docx2txt
import docx
import os
import glob
import subprocess
import re

from os import listdir
from os.path import isfile, join

#argumentos 
#fileDir = r"/Users/wpessoa/repositorios/Fluxo_Editorial/words/"
#fileExt = r".docx"

def get_info(fileDir,files_array):
    while True:
        try:
            #print(os.path.exists(fileDir, fileDir)) 
            #files_array = [_ for _ in os.listdir(fileDir) if _.endswith(fileExt)]
            indice = 0
            image_array=[]
            table_chunks = []
            table_return = []
            
            for file in files_array:
                #
                # raspagem de numero de página do arquivo word ( x.group()) Dados dentro do arquivo app.xml
                # 
                
                pagina_xml_code = subprocess.Popen(["unzip", "-p", fileDir+file , "docProps/app.xml"], stdout=subprocess.PIPE)
                output = pagina_xml_code.communicate()[0]
                #print(output)
                pagina_xml_str = output.decode("utf-8")
                x = re.search('(?<=\<Pages\>)(.*)(?=\<\/Pages\>)', pagina_xml_str)
                # print(file)
                try:
                    pagina_xml = x.group()
                except ValueError:
                    print("Oops! XML não encontrado") 
                #print(pagina_xml)
                
                
                # abrir conecção com o arquivo Word 
                #doc = docx.Document(fileDir+'/'+file)
                doc = docx.Document(fileDir+file)


                # data criação
                dt_doc = doc.core_properties.created
                # ler em cada parágrafo dentro do arquivo Word
                paragra= [p.text for p in doc.paragraphs]

                # extai a quantidade de caracteres
                #caracteres = docx2txt.process(fileDir+'/'+file)
                content = docx2txt.process(fileDir+file)
                #### novo codigo
                caracteres = []
                for line in content.splitlines():
                    #This will ignore empty/blank lines. 
                    if line != '':
                        #Append to list
                        caracteres.append(len(line))
                
                # fim novo codigo
                #### novo codigo
                # pegando quantidade de estilos que o docuento word possui 
                # colocando na variável (array_styles)
                from docx.enum.style import WD_STYLE_TYPE
                styles = doc.styles
                array_styles=[]
                paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
                for style in paragraph_styles:
                    array_styles.append(style.name)
                    #print(style.name)
                #print(len(array_styles))
                # fim novo codigo

                # gerar um array com tamanho de todas as imagens
                for image in doc.inline_shapes:
                    image_array.append([image.width, image.height])
                    # print (image.width, image.height)

                # gerar um array com todas as tabelas
                for table in doc.tables:
                    table_chunks.append(table)
                    
                # gerar lista resposta
                list_retur=file[2:8],pagina_xml,sum(caracteres), len(table_chunks), len(image_array),len(array_styles), dt_doc.date()
                
                # gerar tabela resposta
                table_return.append(list_retur)
            #return(files_array[indice][2:8],len(caracteres), len(table_chunks), len(image_array),dt_doc.date())
            return table_return
            #break
        except ValueError:
            print("Oops!  erro no método verificar diret'ório e extensão do arquivos no argumento desse método")
           

    

